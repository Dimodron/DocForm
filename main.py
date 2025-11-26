import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
import copy
import shutil
import re

import pdfplumber
from docx import Document
import pandas as pd

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STORAGE_DIR = os.path.join(BASE_DIR, "storage")
os.makedirs(STORAGE_DIR, exist_ok=True)

CONFIG_FILE = os.path.join(STORAGE_DIR, "fields_config.json")

DEFAULT_PROFILE_NAME = "default"

DEFAULT_FIELDS = [
    {
        "name": "NAME",
        "label": "ФИО",
        "type": "text",
    },
    {
        "name": "ORGANIZATION",
        "label": "Организация",
        "type": "text",
    },
    {
        "name": "COMMENT",
        "label": "Комментарий",
        "type": "multiline",
    },
]


def default_config():
    return {
        "current_profile": DEFAULT_PROFILE_NAME,
        "profiles": {
            DEFAULT_PROFILE_NAME: {
                "fields": copy.deepcopy(DEFAULT_FIELDS),
                "template_path": None,
            }
        },
    }


def load_config():
    if not os.path.exists(CONFIG_FILE):
        return default_config()

    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return default_config()

    if isinstance(data, list):
        return {
            "current_profile": DEFAULT_PROFILE_NAME,
            "profiles": {
                DEFAULT_PROFILE_NAME: {
                    "fields": data,
                    "template_path": None,
                }
            },
        }

    if isinstance(data, dict) and "fields" in data and "profiles" not in data:
        return {
            "current_profile": DEFAULT_PROFILE_NAME,
            "profiles": {
                DEFAULT_PROFILE_NAME: {
                    "fields": data.get("fields", copy.deepcopy(DEFAULT_FIELDS)),
                    # тут может быть абсолютный путь старого формата
                    "template_path": data.get("template_path"),
                }
            },
        }

    if isinstance(data, dict) and "profiles" in data:
        profiles = data.get("profiles") or {}
        if not profiles:
            return default_config()

        current = data.get("current_profile") or list(profiles.keys())[0]
        if current not in profiles:
            current = list(profiles.keys())[0]

        for name, prof in list(profiles.items()):
            if not isinstance(prof, dict):
                profiles[name] = {
                    "fields": copy.deepcopy(DEFAULT_FIELDS),
                    "template_path": None,
                }
                continue
            if "fields" not in prof:
                prof["fields"] = copy.deepcopy(DEFAULT_FIELDS)
            if "template_path" not in prof:
                prof["template_path"] = None

        return {
            "current_profile": current,
            "profiles": profiles,
        }

    return default_config()


def save_config(profiles, current_profile):
    try:
        data = {
            "current_profile": current_profile,
            "profiles": profiles,
        }
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print("Не удалось сохранить конфиг:", e)


class FileFormApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("DocForm")
        self.geometry("1300x750")

        self.current_file_path = None

        cfg = load_config()
        self.profiles = cfg["profiles"]
        self.current_profile = cfg["current_profile"]

        self.field_widgets = {}
        self.checkbox_vars = {}
        self._build_ui()
        self._load_profile_into_ui()

    def _get_profile(self, name=None):
        if name is None:
            name = self.current_profile
        return self.profiles[name]

    def _save_all_config(self):
        save_config(self.profiles, self.current_profile)

    def _get_template_abs_path(self, template_path: str | None) -> str | None:
        if not template_path:
            return None
        if os.path.isabs(template_path):
            return template_path
        return os.path.join(STORAGE_DIR, template_path)

    def _build_ui(self):
        main = ttk.Frame(self)
        main.pack(fill="both", expand=True, padx=10, pady=10)

        left = ttk.Frame(main)
        right = ttk.Frame(main)

        left.pack(side="left", fill="both", expand=True)
        right.pack(side="right", fill="y")

        top_left = ttk.Frame(left)
        top_left.pack(fill="x")

        open_btn = ttk.Button(top_left, text="Открыть файл", command=self.open_file)
        open_btn.pack(side="left")

        self.file_label = ttk.Label(top_left, text="Файл не выбран")
        self.file_label.pack(side="left", padx=10)

        text_frame = ttk.Frame(left)
        text_frame.pack(fill="both", expand=True, pady=(10, 0))

        self.text = tk.Text(text_frame, wrap="word")
        scroll = ttk.Scrollbar(text_frame, command=self.text.yview)
        self.text.configure(yscrollcommand=scroll.set)

        self.text.pack(side="left", fill="both", expand=True)
        scroll.pack(side="right", fill="y")

        self._init_text_context_menu()

        profiles_frame = ttk.Frame(right)
        profiles_frame.pack(fill="x", pady=(0, 5))

        ttk.Label(profiles_frame, text="Шаблон полей:").pack(side="left")

        self.profile_combo = ttk.Combobox(
            profiles_frame,
            state="readonly",
            values=list(self.profiles.keys()),
            width=25,
        )
        self.profile_combo.pack(side="left", padx=5)
        self.profile_combo.bind("<<ComboboxSelected>>", self.on_profile_change)

        new_profile_btn = ttk.Button(
            profiles_frame, text="Новый шаблон", command=self.create_profile_from_current
        )
        new_profile_btn.pack(side="left", padx=(5, 0))

        delete_profile_btn = ttk.Button(
            profiles_frame, text="Удалить шаблон", command=self.delete_current_profile
        )
        delete_profile_btn.pack(side="left", padx=(5, 0))

        right_title_frame = ttk.Frame(right)
        right_title_frame.pack(fill="x")

        right_title = ttk.Label(
            right_title_frame, text="Поля для заполнения", font=("Arial", 12, "bold")
        )
        right_title.pack(side="left", pady=(10, 10))

        manage_btn = ttk.Button(
            right_title_frame, text="Управлять полями", command=self.open_manage_fields_dialog
        )
        manage_btn.pack(side="right")

        template_btn = ttk.Button(
            right_title_frame, text="DOCX шаблон", command=self.choose_template
        )
        template_btn.pack(side="right", padx=(5, 10))

        self.form_frame = ttk.Frame(right)
        self.form_frame.pack(fill="both", expand=True, pady=(5, 0))

        btn_frame = ttk.Frame(right)
        btn_frame.pack(fill="x", pady=10)

        report_btn = ttk.Button(btn_frame, text="Сохранить отчёт", command=self.save_report)
        report_btn.pack(fill="x")

    def _load_profile_into_ui(self):
        profile = self._get_profile()
        self.fields = profile["fields"]
        self.template_path = profile["template_path"]

        self.profile_combo["values"] = list(self.profiles.keys())
        self.profile_combo.set(self.current_profile)

        self.build_form()

    def _init_text_context_menu(self):
        self.text_menu = tk.Menu(self, tearoff=0)
        self.text_menu.add_command(label="Копировать", command=self.copy_selection_text)
        self.text_menu.add_command(label="Вставить", command=self.paste_into_text)

        self.text.bind("<Button-3>", self._show_text_menu)
        self.text.bind("<Control-c>", self._on_ctrl_c_text)
        self.text.bind("<Control-v>", self._on_ctrl_v_text)

    def _show_text_menu(self, event):
        try:
            self.text_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.text_menu.grab_release()

    def copy_selection_text(self, event=None):
        try:
            selection = self.text.get("sel.first", "sel.last")
        except tk.TclError:
            return
        self.clipboard_clear()
        self.clipboard_append(selection)

    def paste_into_text(self, event=None):
        try:
            data = self.clipboard_get()
        except tk.TclError:
            return
        self.text.insert(tk.INSERT, data)

    def _on_ctrl_c_text(self, event):
        self.copy_selection_text()
        return "break"

    def _on_ctrl_v_text(self, event):
        self.paste_into_text()
        return "break"

    def _attach_entry_context_menu(self, widget: tk.Widget):
        menu = tk.Menu(widget, tearoff=0)
        menu.add_command(label="Вырезать", command=lambda: self._entry_cut(widget))
        menu.add_command(label="Копировать", command=lambda: self._entry_copy(widget))
        menu.add_command(label="Вставить", command=lambda: self._entry_paste(widget))

        def show_menu(event):
            try:
                menu.tk_popup(event.x_root, event.y_root)
            finally:
                menu.grab_release()

        widget.bind("<Button-3>", show_menu)
        widget.bind("<Control-c>", lambda e: (self._entry_copy(widget), "break"))
        widget.bind("<Control-v>", lambda e: (self._entry_paste(widget), "break"))
        widget.bind("<Control-x>", lambda e: (self._entry_cut(widget), "break"))

    def _entry_copy(self, widget):
        try:
            selection = widget.selection_get()
        except tk.TclError:
            return
        self.clipboard_clear()
        self.clipboard_append(selection)

    def _entry_cut(self, widget):
        try:
            selection = widget.selection_get()
        except tk.TclError:
            return
        self.clipboard_clear()
        self.clipboard_append(selection)
        widget.delete("sel.first", "sel.last")

    def _entry_paste(self, widget):
        try:
            data = self.clipboard_get()
        except tk.TclError:
            return
        widget.insert(tk.INSERT, data)

    def _attach_text_context_menu(self, widget: tk.Text):
        menu = tk.Menu(widget, tearoff=0)
        menu.add_command(label="Вырезать", command=lambda: self._text_cut(widget))
        menu.add_command(label="Копировать", command=lambda: self._text_copy(widget))
        menu.add_command(label="Вставить", command=lambda: self._text_paste(widget))

        def show_menu(event):
            try:
                menu.tk_popup(event.x_root, event.y_root)
            finally:
                menu.grab_release()

        widget.bind("<Button-3>", show_menu)
        widget.bind("<Control-c>", lambda e: (self._text_copy(widget), "break"))
        widget.bind("<Control-v>", lambda e: (self._text_paste(widget), "break"))
        widget.bind("<Control-x>", lambda e: (self._text_cut(widget), "break"))

    def _text_copy(self, widget):
        try:
            selection = widget.get("sel.first", "sel.last")
        except tk.TclError:
            return
        self.clipboard_clear()
        self.clipboard_append(selection)

    def _text_cut(self, widget):
        try:
            selection = widget.get("sel.first", "sel.last")
        except tk.TclError:
            return
        self.clipboard_clear()
        self.clipboard_append(selection)
        widget.delete("sel.first", "sel.last")

    def _text_paste(self, widget):
        try:
            data = self.clipboard_get()
        except tk.TclError:
            return
        widget.insert(tk.INSERT, data)

    def build_form(self):
        for child in self.form_frame.winfo_children():
            child.destroy()
        self.field_widgets.clear()
        self.checkbox_vars.clear()

        for i, f in enumerate(self.fields):
            label = ttk.Label(
                self.form_frame, text=f.get("label", f.get("name", "?"))
            )
            label.grid(row=i * 2, column=0, sticky="w", pady=(0, 2))

            ftype = f.get("type", "text")

            if ftype == "multiline":
                widget = tk.Text(self.form_frame, height=4, width=40, wrap="word")
                widget.grid(row=i * 2 + 1, column=0, sticky="we", pady=(0, 8))
                self._attach_text_context_menu(widget)
            elif ftype == "checkbox":
                var = tk.BooleanVar(value=False)
                widget = ttk.Checkbutton(self.form_frame, variable=var)
                widget.grid(row=i * 2 + 1, column=0, sticky="w", pady=(0, 8))
                self.checkbox_vars[f["name"]] = var
            else:
                widget = ttk.Entry(self.form_frame, width=40)
                widget.grid(row=i * 2 + 1, column=0, sticky="we", pady=(0, 8))
                self._attach_entry_context_menu(widget)

            self.field_widgets[f["name"]] = widget

        self.form_frame.columnconfigure(0, weight=1)

    def on_profile_change(self, event=None):
        new_profile = self.profile_combo.get()
        if new_profile == self.current_profile:
            return

        self._get_profile()["fields"] = self.fields
        self._get_profile()["template_path"] = self.template_path

        self.current_profile = new_profile
        self._save_all_config()
        self._load_profile_into_ui()

    def create_profile_from_current(self):
        from tkinter import simpledialog

        name = simpledialog.askstring(
            "Новый шаблон полей",
            "Введите название шаблона:",
            parent=self,
        )
        if not name:
            return
        name = name.strip()
        if not name:
            return
        if name in self.profiles:
            messagebox.showerror("Ошибка", "Шаблон с таким названием уже существует.")
            return

        self._get_profile()["fields"] = self.fields
        self._get_profile()["template_path"] = self.template_path

        self.profiles[name] = {
            "fields": copy.deepcopy(self.fields),
            "template_path": self.template_path,
        }
        self.current_profile = name
        self._save_all_config()
        self._load_profile_into_ui()
        messagebox.showinfo("Создано", f"Создан новый шаблон полей: {name}")

    def delete_current_profile(self):
        if len(self.profiles) <= 1:
            messagebox.showerror(
                "Ошибка",
                "Невозможно удалить единственный шаблон. Должен остаться хотя бы один.",
            )
            return

        if not messagebox.askyesno(
            "Удалить шаблон",
            f"Удалить шаблон полей '{self.current_profile}'?",
            parent=self,
        ):
            return

        prof = self._get_profile()
        tmpl_abs = self._get_template_abs_path(prof.get("template_path"))

        if tmpl_abs and os.path.commonpath([tmpl_abs, STORAGE_DIR]) == STORAGE_DIR:
            if os.path.exists(tmpl_abs):
                try:
                    os.remove(tmpl_abs)
                except OSError:
                    pass

        del self.profiles[self.current_profile]

        self.current_profile = list(self.profiles.keys())[0]
        self._save_all_config()
        self._load_profile_into_ui()
        messagebox.showinfo("Удалено", "Шаблон полей удалён.")

    def open_manage_fields_dialog(self):
        dialog = tk.Toplevel(self)
        dialog.title(f"Управление полями ({self.current_profile})")
        dialog.grab_set()
        dialog.resizable(False, False)

        ttk.Label(dialog, text="Поля (в текущем порядке):").grid(
            row=0, column=0, columnspan=3, sticky="w", padx=10, pady=(10, 5)
        )

        listbox = tk.Listbox(dialog, height=min(15, len(self.fields) or 5), width=50)
        listbox.grid(row=1, column=0, columnspan=3, padx=10, sticky="we")

        def refresh_listbox(selected_index=None):
            listbox.delete(0, tk.END)
            for f in self.fields:
                label = f.get("label", f.get("name", "?"))
                name = f.get("name", "")
                ftype = f.get("type", "text")
                listbox.insert(tk.END, f"{label} ({name}, {ftype})")
            if selected_index is not None and 0 <= selected_index < listbox.size():
                listbox.selection_set(selected_index)
                listbox.activate(selected_index)

        refresh_listbox()

        btn_frame = ttk.Frame(dialog)
        btn_frame.grid(row=2, column=0, columnspan=3, sticky="e", padx=10, pady=10)

        def move_up():
            sel = listbox.curselection()
            if not sel:
                return
            i = sel[0]
            if i == 0:
                return
            self.fields[i - 1], self.fields[i] = self.fields[i], self.fields[i - 1]
            self._get_profile()["fields"] = self.fields
            self._save_all_config()
            self.build_form()
            refresh_listbox(i - 1)

        def move_down():
            sel = listbox.curselection()
            if not sel:
                return
            i = sel[0]
            if i >= len(self.fields) - 1:
                return
            self.fields[i + 1], self.fields[i] = self.fields[i], self.fields[i + 1]
            self._get_profile()["fields"] = self.fields
            self._save_all_config()
            self.build_form()
            refresh_listbox(i + 1)

        def add_field():
            field = self._field_edit_dialog(dialog, None)
            if field is None:
                return
            self.fields.append(field)
            self._get_profile()["fields"] = self.fields
            self._save_all_config()
            self.build_form()
            refresh_listbox(len(self.fields) - 1)

        def edit_field():
            sel = listbox.curselection()
            if not sel:
                messagebox.showerror("Ошибка", "Выберите поле для редактирования.")
                return
            i = sel[0]
            field = self.fields[i]
            updated = self._field_edit_dialog(dialog, field)
            if updated is None:
                return
            self.fields[i] = updated
            self._get_profile()["fields"] = self.fields
            self._save_all_config()
            self.build_form()
            refresh_listbox(i)

        def delete_field():
            sel = listbox.curselection()
            if not sel:
                messagebox.showerror("Ошибка", "Выберите поле для удаления.")
                return
            i = sel[0]
            field = self.fields[i]
            if messagebox.askyesno(
                "Подтверждение",
                f"Удалить поле '{field.get('label', field.get('name', '?'))}'?",
                parent=dialog,
            ):
                del self.fields[i]
                self._get_profile()["fields"] = self.fields
                self._save_all_config()
                self.build_form()
                refresh_listbox(min(i, len(self.fields) - 1))

        up_btn = ttk.Button(btn_frame, text="↑", width=3, command=move_up)
        up_btn.pack(side="left", padx=(0, 5))

        down_btn = ttk.Button(btn_frame, text="↓", width=3, command=move_down)
        down_btn.pack(side="left", padx=(0, 15))

        add_btn = ttk.Button(btn_frame, text="Добавить", command=add_field)
        add_btn.pack(side="left")

        edit_btn = ttk.Button(btn_frame, text="Изменить", command=edit_field)
        edit_btn.pack(side="left", padx=5)

        del_btn = ttk.Button(btn_frame, text="Удалить", command=delete_field)
        del_btn.pack(side="left")

        close_btn = ttk.Button(btn_frame, text="Закрыть", command=dialog.destroy)
        close_btn.pack(side="left", padx=(10, 0))

    def _field_edit_dialog(self, parent, field):
        is_edit = field is not None

        dialog = tk.Toplevel(parent)
        dialog.title("Изменить поле" if is_edit else "Новое поле")
        dialog.grab_set()
        dialog.resizable(False, False)

        ttk.Label(
            dialog,
            text="Внутреннее имя (используется в DOCX как {{ИМЯ}}):",
        ).grid(row=0, column=0, sticky="w", padx=10, pady=(10, 2))
        name_entry = ttk.Entry(dialog, width=30)
        name_entry.grid(row=1, column=0, sticky="we", padx=10)

        ttk.Label(dialog, text="Метка (что показывать в программе):").grid(
            row=2, column=0, sticky="w", padx=10, pady=(10, 2)
        )
        label_entry = ttk.Entry(dialog, width=30)
        label_entry.grid(row=3, column=0, sticky="we", padx=10)

        ttk.Label(dialog, text="Тип поля:").grid(
            row=4, column=0, sticky="w", padx=10, pady=(10, 2)
        )
        type_combo = ttk.Combobox(dialog, values=["text", "multiline", "checkbox"], state="readonly")
        type_combo.grid(row=5, column=0, sticky="we", padx=10)

        if is_edit:
            name_entry.insert(0, field.get("name", ""))
            label_entry.insert(0, field.get("label", ""))
            ftype = field.get("type", "text")
            if ftype not in ("text", "multiline", "checkbox"):
                ftype = "text"
            type_combo.set(ftype)
        else:
            type_combo.set("text")

        result = {"value": None}

        def on_save():
            name = name_entry.get().strip().upper()
            label = label_entry.get().strip()
            ftype = (type_combo.get() or "text").strip()

            if not name:
                messagebox.showerror("Ошибка", "Имя поля не может быть пустым.", parent=dialog)
                return

            for f in self.fields:
                if f is field:
                    continue
                if f.get("name") == name:
                    messagebox.showerror("Ошибка", "Поле с таким именем уже существует.", parent=dialog)
                    return

            result["value"] = {
                "name": name,
                "label": label or name,
                "type": ftype,
            }
            dialog.destroy()

        btn_frame = ttk.Frame(dialog)
        btn_frame.grid(row=6, column=0, sticky="e", padx=10, pady=10)

        save_btn = ttk.Button(btn_frame, text="Сохранить", command=on_save)
        save_btn.pack(side="right", padx=(5, 0))

        cancel_btn = ttk.Button(btn_frame, text="Отмена", command=dialog.destroy)
        cancel_btn.pack(side="right")

        dialog.columnconfigure(0, weight=1)
        dialog.wait_window()

        return result["value"]

    def choose_template(self):
        path = filedialog.askopenfilename(
            title="Выберите корпоративный шаблон (.docx)",
            filetypes=(("Документ Word", "*.docx"), ("Все файлы", "*.*")),
        )
        if not path:
            return

        old_rel = self.template_path
        old_abs = self._get_template_abs_path(old_rel)
        if old_abs and os.path.commonpath([old_abs, STORAGE_DIR]) == STORAGE_DIR:
            if os.path.exists(old_abs):
                try:
                    os.remove(old_abs)
                except OSError:
                    pass

        safe_profile = re.sub(r"[^A-Za-z0-9_-]+", "_", self.current_profile)
        new_rel_name = f"{safe_profile}.docx"
        new_abs = os.path.join(STORAGE_DIR, new_rel_name)

        try:
            shutil.copy2(path, new_abs)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось скопировать шаблон:\n{e}")
            return

        self.template_path = new_rel_name
        self._get_profile()["template_path"] = new_rel_name
        self._save_all_config()
        messagebox.showinfo("Шаблон установлен", f"Будет использоваться шаблон:\n{new_rel_name}")

    def open_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите файл",
            filetypes=(
                ("Все поддерживаемые", "*.txt *.log *.md *.pdf *.doc *.docx *.xls *.xlsx"),
                ("PDF", "*.pdf"),
                ("Word", "*.doc *.docx"),
                ("Excel", "*.xls *.xlsx"),
                ("Текстовые", "*.txt *.log *.md"),
                ("Все файлы", "*.*"),
            ),
        )
        if not file_path:
            return

        self.current_file_path = file_path
        self.file_label.config(text=file_path)

        self.text.config(state="normal")
        self.text.delete("1.0", tk.END)

        try:
            content = self.read_any_file(file_path)
            self.text.insert(tk.END, content)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл:\n{e}")

    def read_any_file(self, path: str) -> str:
        path_l = path.lower()
        if path_l.endswith((".txt", ".log", ".md")):
            return self.read_text(path)
        if path_l.endswith(".pdf"):
            return self.read_pdf(path)
        if path_l.endswith((".doc", ".docx")):
            return self.read_word(path)
        if path_l.endswith((".xls", ".xlsx")):
            return self.read_excel(path)
        return "Формат файла не поддерживается."

    @staticmethod
    def read_text(path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def read_pdf(path: str) -> str:
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return text or "PDF не содержит распознаваемый текст (возможно, только картинки)."

    @staticmethod
    def read_word(path: str) -> str:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def read_excel(path: str) -> str:
        df_dict = pd.read_excel(path, sheet_name=None)
        output = ""
        for sheet_name, df in df_dict.items():
            output += f"=== Лист: {sheet_name} ===\n"
            output += df.to_string(index=False)
            output += "\n\n"
        return output

    def collect_form_data(self) -> dict:
        data = {}
        for f in self.fields:
            name = f.get("name")
            ftype = f.get("type", "text")
            widget = self.field_widgets.get(name)
            if not widget:
                continue

            if ftype == "multiline":
                value = widget.get("1.0", tk.END).strip()
            elif ftype == "checkbox":
                var = self.checkbox_vars.get(name)
                value = bool(var.get()) if var is not None else False
            else:
                value = widget.get().strip()
            data[name] = value
        return data

    @staticmethod
    def _sanitize_filename(name: str) -> str:
        for ch in '<>:"/\\|?*':
            name = name.replace(ch, "_")
        name = name.strip().replace(" ", "_")
        return name or "report"

    @staticmethod
    def _apply_template(doc: Document, placeholders: dict[str, str]):

        def process_paragraph(paragraph):
            if not paragraph.runs:
                return

            full_text = "".join(run.text for run in paragraph.runs)
            new_text = full_text
            for ph, val in placeholders.items():
                new_text = new_text.replace(ph, val)

            if new_text == full_text:
                return

            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

        for p in doc.paragraphs:
            process_paragraph(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p)

        for section in doc.sections:
            header = section.header
            footer = section.footer
            for p in header.paragraphs:
                process_paragraph(p)
            for p in footer.paragraphs:
                process_paragraph(p)

    def save_report(self):
        data = self.collect_form_data()

        default_name = "report.docx"
        full_name_value = data.get("NAME")
        if isinstance(full_name_value, str):
            full_name_value = full_name_value.strip()
        if full_name_value:
            base = self._sanitize_filename(full_name_value)
            default_name = f"{base}.docx"
        else:
            for f in self.fields:
                ftype = f.get("type", "text")
                if ftype not in ("text", "multiline"):
                    continue
                val = data.get(f.get("name"))
                if isinstance(val, str):
                    val = val.strip()
                if val:
                    base = self._sanitize_filename(val)
                    default_name = f"{base}.docx"
                    break

        save_path = filedialog.asksaveasfilename(
            title="Сохранить отчёт",
            defaultextension=".docx",
            filetypes=(("Документ Word", "*.docx"),),
            initialfile=default_name,
        )
        if not save_path:
            return

        try:
            tmpl_abs = self._get_template_abs_path(self.template_path)
            if tmpl_abs and os.path.exists(tmpl_abs):
                doc = Document(tmpl_abs)

                placeholders = {}
                for f in self.fields:
                    name = f.get("name")
                    ftype = f.get("type", "text")
                    raw_value = data.get(name)

                    if ftype == "checkbox":
                        val = "Да" if raw_value else ""
                    else:
                        val = (raw_value or "").strip()

                    ph = f"{{{{{name}}}}}"  # {{NAME}}
                    placeholders[ph] = val

                self._apply_template(doc, placeholders)
            else:
                doc = Document()
                for f in self.fields:
                    name = f.get("name")
                    label = f.get("label", name)
                    ftype = f.get("type", "text")
                    raw_value = data.get(name)

                    if ftype == "checkbox":
                        if not raw_value:
                            continue
                        val = "Да"
                    else:
                        val = (raw_value or "").strip()
                        if not val:
                            continue

                    doc.add_paragraph(f"{label}: {val}")

            doc.save(save_path)
            messagebox.showinfo("Готово", f"Отчёт сохранён:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить отчёт:\n{e}")


if __name__ == "__main__":
    app = FileFormApp()
    app.mainloop()
